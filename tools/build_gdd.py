from __future__ import annotations

import json
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
BALANCE_PATH = DOCS / "balance_results_v1.1.json"
OUTPUT = DOCS / "NEXUS_2040_Ghost_Grid_GDD_v1.1.docx"

BLUE = "2E74B5"
DARK_BLUE = "0B2545"
MID_BLUE = "1F4D78"
GREEN = "008F5D"
NEON_GREEN = "00A86B"
CYAN = "1976A3"
GOLD = "7A5A00"
RED = "9B1C1C"
INK = "243447"
MUTED = "5B6775"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E9F7F1"
LIGHT_GOLD = "FFF7DD"
LIGHT_RED = "FCEBEC"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"
GRID = "C8D2DE"

PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(
    run,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name: str, size: float, color: str = INK, bold: bool | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if bold is not None:
        style.font.bold = bold


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_left_border(paragraph, color: str, size: str = "18", space: str = "8") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = pbdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        pbdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)


def add_field(paragraph, instruction: str, display: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Calibri", 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    set_style_font(title, "Calibri", 28, DARK_BLUE, True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(7)
    title.paragraph_format.keep_with_next = True

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, "Calibri", 14, MID_BLUE)
    subtitle.font.italic = False
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle.paragraph_format.keep_with_next = True

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, MID_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        set_style_font(style, "Calibri", size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    set_style_font(caption, "Calibri", 9.5, MUTED)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.keep_with_next = False

    for name, font, size, color, after, line in [
        ("Formula", "Consolas", 9.5, DARK_BLUE, 7, 1.15),
        ("Callout", "Calibri", 10.5, INK, 8, 1.18),
        ("Small Text", "Calibri", 9, MUTED, 4, 1.1),
        ("TOC Entry", "Calibri", 10.5, INK, 3, 1.1),
    ]:
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style = styles[name]
        set_style_font(style, font, size, color)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def setup_running_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = ""
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        left = p.add_run("NEXUS 2040: GHOST GRID")
        set_run_font(left, size=8.5, color=MUTED, bold=True)
        p.add_run("\t")
        right = p.add_run("GDD v1.1 | 9 Ağustos 2026")
        set_run_font(right, size=8.5, color=MUTED)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = fp.add_run("Sayfa ")
        set_run_font(run, size=8.5, color=MUTED)
        add_field(fp, "PAGE", "1")


def create_numbering(doc: Document, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    lvl.append(ppr)
    if not ordered:
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), "Segoe UI Symbol")
        rfonts.set(qn("w:hAnsi"), "Segoe UI Symbol")
        rpr.append(rfonts)
        lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_list(doc: Document, items: Iterable[str], ordered: bool = False) -> None:
    num_id = create_numbering(doc, ordered)
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, num_id)
        run = p.add_run(item)
        set_run_font(run, size=11, color=INK)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        prefix = p.add_run(bold_prefix)
        set_run_font(prefix, bold=True, color=INK)
        rest = p.add_run(text[len(bold_prefix) :])
        set_run_font(rest, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, color=INK)


def add_callout(doc: Document, label: str, text: str, kind: str = "info") -> None:
    palette = {
        "info": (LIGHT_BLUE, BLUE),
        "positive": (LIGHT_GREEN, GREEN),
        "caution": (LIGHT_GOLD, GOLD),
        "risk": (LIGHT_RED, RED),
    }
    fill, accent = palette[kind]
    p = doc.add_paragraph(style="Callout")
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.05)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, fill)
    set_paragraph_left_border(p, accent)
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, size=10.5, color=accent, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def add_formula(doc: Document, formula: str) -> None:
    p = doc.add_paragraph(style="Formula")
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    set_paragraph_shading(p, LIGHT_GRAY)
    run = p.add_run(formula)
    set_run_font(run, name="Consolas", size=9.5, color=DARK_BLUE)


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for edge, value in CELL_MARGINS.items():
        node = mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), GRID)
        node.set(qn("w:space"), "0")


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    assert sum(widths_dxa) == CONTENT_WIDTH_DXA, widths_dxa
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblind.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trpr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths_dxa[idx]))
            tcw.set(qn("w:type"), "dxa")
    set_table_borders(table)


def set_repeat_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    widths_dxa: Sequence[int],
    *,
    font_size: float = 9.2,
    center_cols: set[int] | None = None,
) -> None:
    center_cols = center_cols or set()
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(header)
        set_cell_shading(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            for run in p.runs:
                set_run_font(run, size=font_size, color=DARK_BLUE, bold=True)
    set_repeat_header(table.rows[0])
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            if row_idx % 2 == 1:
                set_cell_shading(cells[idx], "F8FAFC")
            for p in cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in center_cols else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.08
                for run in p.runs:
                    set_run_font(run, size=font_size, color=INK)
    set_table_geometry(table, widths_dxa)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(1)


def add_figure(doc: Document, filename: str, caption: str, alt: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    shape = run.add_picture(str(ASSETS / filename), width=Inches(6.25))
    shape._inline.docPr.set("descr", alt)
    c = doc.add_paragraph(caption, style="Caption")
    c.paragraph_format.keep_with_next = False


def add_toc(doc: Document) -> None:
    doc.add_heading("İçindekiler", level=1)
    p = doc.add_paragraph(style="TOC Entry")
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "İçindekiler Word açıldığında güncellenecektir.")
    add_callout(
        doc,
        "Kullanım",
        "Bu belge kalıcı teknoloji stratejisi ile taktik PvE siber operasyonlarını tek ürün sisteminde birleştirir. Sayılar nihai gerçekler değil; simülasyon, telemetri ve kontrollü oyuncu testleriyle sürümlenecek başlangıç hipotezleridir.",
        "info",
    )


def add_cover(doc: Document) -> None:
    # Header pattern: editorial_cover. Named overrides: centered 30 pt title,
    # restrained green kicker and 15 pt subtitle; all other geometry follows
    # compact_reference_guide.
    for _ in range(6):
        doc.add_paragraph().paragraph_format.space_after = Pt(11)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    kr = kicker.add_run("YAKIN GELECEK TEKNOLOJİ STRATEJİSİ")
    set_run_font(kr, size=10, color=GREEN, bold=True)
    title = doc.add_paragraph("NEXUS 2040: GHOST GRID", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_run_font(run, size=30, color=DARK_BLUE, bold=True)
    subtitle = doc.add_paragraph(
        "Kalıcı Teknoloji Konsorsiyumu ve PvE Siber Operasyon Simülasyonu",
        style="Subtitle",
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)

    add_callout(
        doc,
        "Ürün tezi",
        "Stratejik katman operasyona neden verir; operasyon stratejik kararın sonucunu belirler. Oyuncu kalıcı merkezini teknoloji ve ekonomiyle büyütürken, gerçek dünyaya uygulanabilir saldırı tarifi içermeyen taktik görevlerde karar baskısını yaşar.",
        "positive",
    )
    metadata = [
        ("Belge sürümü", "1.1 - birleşik konsept ve üretim temeli"),
        ("Tarih", "9 Ağustos 2026"),
        ("İlk ürün", "Masaüstü tarayıcı, tek oyunculu PvE, yaşayan NPC ekonomili kontrollü dikey kesit"),
        ("Oyuncu rolü", "Bağımsız teknoloji ve siber güvenlik konsorsiyumu yöneticisi"),
        ("Denge kanıtı", "100.000 koşuluk operasyon Monte Carlo + 120 günlük ekonomi stres testi"),
        ("Tasarım duruşu", "Adil, açıklanabilir, ana merkezi güvenli, FOMO'suz ve güvenli gerçekçilik"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=10.5, color=DARK_BLUE, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=10.5, color=INK)
    doc.add_page_break()


def add_executive_summary(doc: Document) -> None:
    doc.add_heading("1. Yönetici özeti ve kilit kararlar", level=1)
    add_body(
        doc,
        "NEXUS 2040: Ghost Grid, oyuncunun 2040 yılında bağımsız bir teknoloji ve siber güvenlik konsorsiyumunu yönettiği, yaşayan NPC dünyalı bir strateji-simülasyon oyunudur. Oyuncu üretim tesisleri kurar, araştırma ve sözleşme portföyünü yönetir; stratejik hedefler taktik PvE siber operasyonlarında çözülür. Kalıcı ana merkez saldırıya kapalıdır ve başarısızlık aylarca emeği silmez.",
    )
    add_body(
        doc,
        "Oyun gerçekçi terminolojiyi, belirsizliği, kaynak kısıtını, tespit ve toparlanma mantığını korur; gerçek hedeflere uygulanabilecek komutları, açıkları, IP adreslerini veya zararlı yazılım davranışlarını içermez. Gerçekçilik teknik icraattan değil, sistemlerin birbirine verdiği nedensel tepkiden doğar.",
    )
    add_table(
        doc,
        ["Karar alanı", "v1.1 kararı"],
        [
            ("Platform", "Masaüstü öncelikli web; 1280x720 alt sınır, 1440x900 tasarım hedefi"),
            ("Oyun modu", "Tek oyunculu PvE; kalıcı profil, NPC pazarı, sözleşmeler ve uyarlanan düşman doktrinleri"),
            ("Çekirdek strateji", "Dünya olayı -> kaynak dağıtımı -> sözleşme -> hazırlık -> operasyon -> dünya sonucu -> yatırım"),
            ("Taktik operasyon", "Keşif -> yükleme -> yürütme -> Node Routing -> çıkış -> adli rapor"),
            ("İlk içerik", "1 bölge, 3 NPC organizasyonu, 5 tesis, 12 hedef, 3 görev tipi, 1 mini oyun"),
            ("Oturum", "3-5 dk yönetim, 12-20 dk operasyon, 45-90 dk kriz zinciri; zorunlu günlük yok"),
            ("Gelir", "Prototipte yok; ileride kozmetik/premium içerik, güç satışı yok"),
            ("Başarı ölçütü", "Oyuncu operasyon sonucunu ve dünyadaki etkisini açıklayabilmeli; ikinci sözleşmeyi gönüllü seçmeli"),
        ],
        [2700, 6660],
        font_size=9.5,
    )
    add_callout(
        doc,
        "En önemli kapsam ilkesi",
        "PvP, klan, oyuncu pazarı, sezon sistemi, mobil istemci, 3D küre ve ek mini oyunlar MVP dışında kalır. İlk prototip yalnız stratejik karar ile operasyon sonucunun birbirini beslediğini kanıtlar.",
        "caution",
    )

    doc.add_heading("1.1 Başarı tanımı", level=2)
    add_list(
        doc,
        [
            "İlk kez oynayan bir kullanıcı 8 dakika içinde hedef seçip ilk operasyonu başlatabilir.",
            "Bir operasyonun bütün aşamaları 8-12 dakikada tamamlanır; oyuncu en az iki anlamlı risk/ödül kararı verir.",
            "Sonuç ekranı başarı, tespit, ödül, ısı, pazar ve NPC ilişki değişimini nedenleriyle açıklar.",
            "Oyuncu kaybettiğinde ilerlemesi silinmez; delilin en az %25'ini analiz verisi olarak kazanır.",
            "Dikey kesit 30 dakikada stratejik hazırlığın taktik sonucu, taktik sonucun da dünyayı değiştirdiğini gösterir.",
            "Günde 2 ve 10 operasyon yapan profiller arasındaki ekonomik ilerleme farkı %20'yi geçmez.",
        ],
    )


def add_research_foundation(doc: Document) -> None:
    doc.add_heading("2. Araştırma temeli ve gerçekçilik sınırı", level=1)
    add_body(
        doc,
        "Tasarım iki araştırma hattını birleştirir. Siber simülasyon tarafı MITRE ATT&CK ve NIST çerçevelerinden yalnız yüksek seviyeli amaç ve savunma mantığını alır. Kalıcı strateji tarafı ise çevrimdışı kayıp baskısını azaltan ve üretim, ekonomi, diplomasi ile farklı oyun tarzlarına alan açan resmî oyun tasarımı örneklerinden yararlanır. [R1-R7, R17-R20]",
    )
    add_table(
        doc,
        ["Araştırma", "Oyuna aktarılan ilke", "MVP karşılığı", "Aktarılmayan içerik"],
        [
            ("MITRE ATT&CK", "Amaç odaklı operasyon evreleri", "Keşif, erişim, hareket, hedef, çıkış", "Gerçek teknik/alt teknik uygulama adımları"),
            ("NIST CSF 2.0", "Riskin yalnız saldırı değil yönetişim ve toparlanma olması", "NPC containment, hizmet etkisi ve adli rapor", "Uyumluluk eğitimi veya gerçek kontrol listesi"),
            ("NIST 800-61r3", "Tespit, müdahale ve toparlanmanın sürekli süreç olması", "NPC alarm ve containment durum makinesi", "Gerçek olay müdahale prosedürü"),
            ("CISA CPG", "Az sayıda yüksek etkili savunma yatırımının önceliği", "Oyuncuya anlaşılır savunma yükseltmeleri", "Sektör bazlı gerçek altyapı ayrıntıları"),
            ("SDT araştırmaları", "Özerklik, yeterlik ve bağ kurma deneyiminin motivasyona katkısı", "Seçilebilir hedefler, açık ustalık, hikâye ekibi", "Zorunlu günlük seri ve kayıp korkusu"),
            ("MDA", "Mekanikten doğan dinamiğin hedeflenen duyguyu üretmesi", "Her özellik için deneyim hipotezi", "Özellik listesinin tek başına 'eğlence' sayılması"),
            ("Hades' Star", "Kalıcı dünya hissi çevrimdışı kaynak hırsızlığı olmadan da kurulabilir", "Güvenli ana merkez; dışarıda sınırlı riskli fırsatlar", "Çevrimdışı baskı ve kalıcı üs kaybı"),
            ("Prosperous Universe", "Üretim zinciri ve şirket yönetimi çatışmasız uzun vadeli strateji üretebilir", "NPC pazarı, uzmanlaşma ve sözleşme portföyü", "Tam oyuncu ekonomisi ve MMO karmaşıklığı"),
            ("Fair Play çerçevesi", "Zarar yalnız moderasyonla değil sistem tasarımıyla da önlenmelidir", "İlk sürüm PvE; gelecekte sınırlı ve isteğe bağlı rekabet", "Tekrarlı hedefleme ve bezdirme döngüsü"),
        ],
        [1500, 2500, 2600, 2760],
        font_size=8.6,
    )

    doc.add_heading("2.1 Güvenli gerçekçilik sözleşmesi", level=2)
    add_table(
        doc,
        ["Gerçekçi tutulacak", "Kurgulanacak / soyutlanacak"],
        [
            ("Kaynak ve zaman bütçesi; düşük gürültü ile hız arasındaki gerilim", "Araç, protokol, servis, ağ adresi ve açık adları"),
            ("Keşif belirsizliği; eksik bilgiyle karar verme", "Topoloji ve hedef verileri; tamamı sunucu üretimli sentetik içerik"),
            ("Tespit, olay müdahalesi, segmentasyon ve toparlanma", "Payload etkileri; kod veya gerçek komut gösterilmez"),
            ("Yetkilendirme, delil bütünlüğü ve operasyon kaydı", "Hikâye kurumları, coğrafya ve para birimleri"),
            ("Savunmanın saldırıyı tamamen önlemekten çok maliyeti artırması", "Fiziksel FLOPS/Gbps değerleri; iç ekonomide normalize değer kullanılır"),
        ],
        [4680, 4680],
        font_size=9.1,
    )
    add_callout(
        doc,
        "İçerik güvenliği",
        "Hiçbir görev gerçek CVE numarası, gerçek kuruluş, kamu IP'si, çalışır terminal komutu, parola kırma tarifi veya zararlı yazılım kodu içermez. Oyunun gerçekçiliği teknik icraattan değil karar sisteminden gelir.",
        "risk",
    )


def add_vision_and_experience(doc: Document) -> None:
    doc.add_heading("3. Vizyon, dünya ve oyuncu deneyimi", level=1)
    doc.add_heading("3.1 Konsept: NEXUS 2040", level=2)
    add_body(
        doc,
        "Yakın gelecekte şehirlerin enerji, lojistik, sağlık, üretim ve iletişim hizmetleri birbirine bağlı otonom ağlarla yönetilir. Oyuncu NODE-7 adlı küçük bir müdahale hücresini devralır ve onu bağımsız teknoloji konsorsiyumuna dönüştürür. Mikro şebeke, veri merkezi, robotik atölye, araştırma laboratuvarı ve güvenlik operasyon merkezi aynı stratejik üs içinde çalışır.",
    )
    add_body(
        doc,
        "Ghost Grid, gerçek altyapının sentetik ikizlerinden oluşan yetkili operasyon ağıdır. Sözleşmeler oyuncuyu bu ağa bağlar; çıkarılan delil, düzeltilen hizmet veya engellenen düşman otomasyonu bölgesel arzı, fiyatları, itibarı ve yeni görevleri değiştirir. Başarı ölçüsü bir sistemi çökertmek değil, amaç için yeterli sonucu minimum iz ve toplumsal maliyetle üretmektir.",
    )

    doc.add_heading("3.2 Oyuncu kimliği ve başarı yolları", level=2)
    add_table(
        doc,
        ["Doktrin", "Stratejik üstünlük", "Operasyon tarzı", "Kampanya hedefi"],
        [
            ("Altyapı Muhafızı", "Enerji ve hizmet dayanıklılığı", "Düşük riskli geri kazanım", "Bölgesel güven"),
            ("İstihbarat Mimarı", "Hedef ve pazar görünürlüğü", "Keşif ve delil kalitesi", "Bilgi üstünlüğü"),
            ("Otomasyon Öncüsü", "Bileşen ve işlem verimliliği", "Hızlı fakat kontrollü rota", "Teknoloji standardı"),
            ("Sözleşme Stratejisti", "Sermaye ve itibar", "Amaca göre değişen dengeli yükleme", "Ekonomik etki"),
        ],
        [1900, 2600, 2600, 2260],
        font_size=8.9,
    )
    add_body(
        doc,
        "Doktrinler sınıf kilidi değildir. Araştırmalar yatay yan seçenekler açar; oyuncu kampanya içinde iki doktrini harmanlayabilir. Zafer tek bir yok etme puanına değil güven, teknoloji, ekonomik etki ve istihbarat hedeflerinden seçilen kombinasyona dayanır.",
    )

    doc.add_heading("3.3 Deneyim sütunları", level=2)
    add_list(
        doc,
        [
            "Yönetici-operatör olma hissi: Oyuncu menülere değil, yaşayan bir teknoloji merkezi ve operasyon odasına hükmettiğini hissetmeli.",
            "Açıklanabilir gerilim: Riskler görünür aralıklarla sunulmalı; kayıp, gizli zar veya hileli zorluk gibi algılanmamalı.",
            "Hazırlık ve ustalık: Doğru istihbarat ve araç eşleşmesi ana belirleyici; mini oyun sonucu destekler ama tek başına belirlemez.",
            "Kayıpsız cesaret: Başarısızlık kalıcı emeği silmez; yeni bilgi, rapor veya daha iyi bir plan üretir.",
            "Dünya etkisi: Operasyon sonucu yalnız ödül ekranında kalmamalı; sözleşme, pazar, ilişki veya kriz durumunu değiştirmeli.",
            "Etik bağlılık: Dönüş nedeni merak, ustalık ve yeni senaryo olmalı; FOMO, çevrimdışı kayıp veya ücretli güç kullanılmamalı.",
        ],
    )

    doc.add_heading("3.4 MDA deneyim haritası", level=2)
    add_table(
        doc,
        ["Hedeflenen estetik", "Mekanik", "Oluşan dinamik"],
        [
            ("Meydan okuma", "Görünür başarı/takip olasılığı, riskli üst kademe", "Oyuncu güvenli hedef ile yüksek ödül arasında karar verir"),
            ("Keşif", "Seed'li ağ topolojileri, eksik istihbarat", "Her hedef önce okunması gereken küçük bir problem olur"),
            ("Ustalık", "Açıklanabilir formüller, tekrar analizi", "Oyuncu şanstan çok planını iyileştirdiğini hisseder"),
            ("Anlatı", "Sektör kampanyaları, adli parçalar", "Başarı sadece sayı değil dünyayı açan delildir"),
            ("İfade", "Araç ve araştırma yüklemeleri", "Aynı hedef sessiz, hızlı veya analitik tarzla çözülebilir"),
            ("Sahiplenme", "Kalıcı güvenli merkez ve uzmanlaşan tesisler", "Oyuncu büyümeyi kayıp korkusuyla değil kimlikle ilişkilendirir"),
        ],
        [1650, 3000, 4710],
        font_size=9.1,
    )


def add_core_loop(doc: Document) -> None:
    doc.add_heading("4. Çekirdek oyun döngüsü ve oturum yapısı", level=1)
    add_callout(
        doc,
        "Birleşim kuralı",
        "Hiçbir taktik operasyon bağlamsız başlamaz; her operasyon bir sözleşmeye, krize veya stratejik hedefe bağlıdır. Hiçbir stratejik karar da yalnız bekleme sayacına dönüşmez; önemli dönüm noktaları oyuncuya oynanabilir karar sunar.",
        "positive",
    )
    add_table(
        doc,
        ["Adım", "Süre", "Oyuncu kararı", "Sistem çıktısı"],
        [
            ("Dünya okuma", "1-2 dk", "Hangi kriz, pazar sinyali veya NPC hamlesi önemli?", "Fırsat, tehdit ve talep değişimi"),
            ("Kaynak dağıtımı", "1-3 dk", "Enerji ve işlem gücü hangi tesise ayrılmalı?", "Üretim ve hazırlık profili"),
            ("Sözleşme", "30-60 sn", "Getiri, itibar veya istihbarat mı?", "Amaç, sınırlı risk ve dünya etkisi"),
            ("Keşif + yükleme", "2-4 dk", "İstihbarat, kapasite ve doktrin nasıl dengelenmeli?", "Başarı bandı, maliyet ve gürültü"),
            ("Operasyon", "5-8 dk", "İlerlemek, rota değiştirmek veya geri çekilmek?", "Erişim, trace ve Node Routing"),
            ("Çıkış", "1-2 dk", "İkincil delil mi, temiz çıkış mı?", "Ödül, tespit ve ısı"),
            ("Dünya sonucu", "1-2 dk", "Delil satılsın, araştırılsın veya ilişki için kullanılsın mı?", "Pazar, NPC ilişkisi ve kriz durumu"),
            ("Yatırım", "2-4 dk", "Tesis, araştırma veya yeni sözleşme hazırlığı?", "Kalıcı fakat yatay ağırlıklı gelişim"),
        ],
        [950, 900, 3400, 4110],
        font_size=8.8,
        center_cols={0, 1},
    )
    add_callout(
        doc,
        "Oturum hedefi",
        "Bir tam operasyon 8-12 dakika, stratejik hazırlık ve sonuçla 12-20 dakika sürer. Yönetim kontrolü 3-5 dakikada anlamlıdır; kriz zinciri 45-90 dakikalık uzun oturuma dönüşebilir. Oyun çıkmayı cezalandırmaz.",
        "positive",
    )

    doc.add_heading("4.1 İlk 30 dakika", level=2)
    add_list(
        doc,
        [
            "0-3 dk: Bölgesel enerji krizi, güvenli ana merkez ve beş kaynak tanıtılır.",
            "3-7 dk: Oyuncu veri merkezine enerji ayırır ve ilk sözleşmesini seçer.",
            "7-12 dk: Kılavuzlu keşif ve yükleme; ekonomik hazırlığın başarı bandını değiştirdiği gösterilir.",
            "12-19 dk: İlk operasyon ve Node Routing; hız/gizlilik kararı verilir.",
            "19-23 dk: Adli rapor; her yüzde, kaynak ve dünya etkisi açıklanır.",
            "23-27 dk: Pazar ve NPC ilişki değişimi görünür; ilk tesis veya araştırma seçilir.",
            "27-30 dk: İkinci sözleşme oyuncuya bırakılır; öğretici geri çekilir.",
        ],
        ordered=True,
    )

    doc.add_heading("4.2 Rutin işlerde manuel oyun zorunluluğu", level=2)
    add_table(
        doc,
        ["Çözüm", "Kullanım", "Getiri", "Tasarım amacı"],
        [
            ("Manuel", "Hikâye, ilk tamamlama, yüksek risk", "%100 ödül ve tam ustalık raporu", "Çekirdek oynanışı yaşatmak"),
            ("Planlı çözüm", "Daha önce ustalaşılan düşük riskli sözleşme", "%78-88 beklenen getiri", "Tekrar yorgunluğunu azaltmak"),
            ("İzleme", "Devam eden üretim ve NPC pazarı", "Karar gerekmedikçe müdahale yok", "Sayaç nöbetini engellemek"),
        ],
        [1500, 3000, 2200, 2660],
        font_size=9.0,
    )


def add_operation_system(doc: Document) -> None:
    doc.add_heading("5. Gerçekçi operasyon simülasyonu", level=1)
    doc.add_heading("5.1 Beş aşamalı operasyon", level=2)
    add_table(
        doc,
        ["Aşama", "Amaç", "Ana oyuncu statı", "NPC savunması", "Kritik karar"],
        [
            ("Keşif", "Hedef yapısını ve belirsizliği azalt", "Keşif + Analiz", "İzleme + Sertleştirme", "Ek kaynak harcayıp tahmini daraltmak"),
            ("Erişim", "Sentetik dış yüzeyde oturum kur", "Erişim + Keşif", "Sertleştirme", "Sessiz veya hızlı araç kullanmak"),
            ("Hareket", "Ağ bölümleri arasında rota bul", "Kontrol + Erişim", "Segmentasyon + İzleme", "Kısa ve riskli ya da uzun ve güvenli rota"),
            ("Hedef", "Delili güvenli biçimde toplamak/servisi düzeltmek", "Kontrol + Analiz", "Dayanıklılık + Segmentasyon", "Asgari hedefte çıkmak veya ikincil delili almak"),
            ("Çıkış", "Bağlantıyı kapat ve trace'i yönet", "Gizlilik + Analiz", "İzleme + Dayanıklılık", "Kazancı korumak veya log temizliği denemek"),
        ],
        [1050, 2500, 1800, 1850, 2160],
        font_size=8.5,
    )

    doc.add_heading("5.2 Oyuncu ve hedef istatistikleri", level=2)
    add_table(
        doc,
        ["Oyuncu statı", "Anlamı", "Hedef karşılığı"],
        [
            ("Keşif", "Belirsizliği azaltma ve uygun giriş profilini bulma", "İzleme / görünürlük"),
            ("Erişim", "İlk oturumu güvenilir biçimde kurma", "Sertleştirme"),
            ("Kontrol", "Ağ içinde görev eylemlerini sürdürme", "Segmentasyon"),
            ("Gizlilik", "Gürültüyü ve tespit olasılığını azaltma", "İzleme"),
            ("Analiz", "Delil kalitesi, çıkış ve sonuç yorumlama", "Dayanıklılık"),
        ],
        [1650, 4600, 3110],
        font_size=9.2,
    )
    add_body(
        doc,
        "Ham statlar 0-100 aralığındadır. 60 sonrasında marjinal fayda azaltılır; böylece ileri oyuncu güçlenir fakat düşük kademeyi tamamen otomatikleştirmez.",
    )
    add_formula(doc, "effective_stat(s) = s, s <= 60; aksi halde 60 + 0.55 x (s - 60)")

    doc.add_heading("5.3 Başarı olasılığı", level=2)
    add_body(
        doc,
        "Her görev tipi oyuncu statlarını ve hedef savunmalarını farklı ağırlıklarla birleştirir. Ağırlıkların her satırı 1.00 eder; böylece bütün görevler aynı 0-100 ölçeğinde karşılaştırılır.",
    )
    add_table(
        doc,
        ["Görev", "Keşif", "Erişim", "Kontrol", "Gizlilik", "Analiz"],
        [
            ("Delil Kurtarma", "0.22", "0.24", "0.20", "0.20", "0.14"),
            ("Hizmet Geri Kazanım", "0.12", "0.18", "0.38", "0.10", "0.22"),
            ("Atıf İzleme", "0.32", "0.16", "0.10", "0.24", "0.18"),
        ],
        [1900, 1492, 1492, 1492, 1492, 1492],
        font_size=8.9,
        center_cols={1, 2, 3, 4, 5},
    )
    add_table(
        doc,
        ["Görev", "Sertleştirme", "İzleme", "Segmentasyon", "Dayanıklılık"],
        [
            ("Delil Kurtarma", "0.28", "0.28", "0.22", "0.22"),
            ("Hizmet Geri Kazanım", "0.20", "0.15", "0.25", "0.40"),
            ("Atıf İzleme", "0.15", "0.40", "0.15", "0.30"),
        ],
        [2100, 1815, 1815, 1815, 1815],
        font_size=8.9,
        center_cols={1, 2, 3, 4},
    )
    add_body(doc, "Aşağıdaki açık form Delil Kurtarma görevinin örneğidir; diğer görevler tabloda verilen ağırlıkları aynı çözümleyiciye geçirir.")
    add_formula(doc, "A = 0.22xKeşif + 0.24xErişim + 0.20xKontrol + 0.20xGizlilik + 0.14xAnaliz")
    add_formula(doc, "D = 0.28xSertleştirme + 0.28xİzleme + 0.22xSegmentasyon + 0.22xDayanıklılık")
    add_formula(doc, "z = (A + araç_puanı - D)/11 + 0.18xln(1+istihbarat) + beceri_bonus - ısı_cezası + destek")
    add_formula(doc, "P_başarı = 0.08 + 0.84 x sigmoid(z)")
    add_body(
        doc,
        "Alt ve üst sınırların %8 ve %92 olması, zayıf oyuncuya küçük bir şans bırakırken hiçbir yapının kesin sonuç vermemesini sağlar. Isı başarıyı yalnızca 40 sonrasında etkiler: ısı_cezası = max(0, Isı-40) x 0.006.",
    )
    add_formula(doc, "başarı = [U0 < P_başarı]; tespit = [U1 < P_tespit]; ödül_jitter = 0.92 + 0.16xU2")
    add_body(doc, "U0, U1 ve U2 operasyon seed'inden üretilen bağımsız [0,1) değerleridir. Böylece sonuç sunucuda tekrar üretilebilir ve aynı rastgele sayı iki ayrı karar için kullanılmaz.")
    add_formula(doc, "tahmin_bant_yarıçapı = max(2, 12 - 2xistihbarat_birimi) yüzde puan")
    add_body(doc, "Planlama ekranı kesin bir yüzde yerine P_başarı +/- tahmin bandını gösterir. İstihbarat yalnız olasılığı artırmaz, oyuncunun belirsizliğini de azaltır.")
    add_figure(
        doc,
        "success_curve_v1_1.png",
        "Şekil 1. Başarı eğrisi; istihbarat ve iyi mini oyun performansı olasılığı kaydırır fakat garanti oluşturmaz.",
        "Saldırı ve savunma puanı farkına göre hazırlıksız ve hazırlıklı operasyonların başarı yüzdesi eğrileri.",
    )

    doc.add_heading("5.4 Tespit ve ısı modeli", level=2)
    add_formula(doc, "skor_gürültüsü = 6 - 0.12 x mini_oyun_skoru")
    add_formula(doc, "q = (İzleme + araç_gürültüsü + skor_gürültüsü + 0.25xIsı - Gizlilik - örtü)/12")
    add_formula(doc, "P_tespit = 0.03 + 0.55 x sigmoid(q)")
    add_body(
        doc,
        "Tespit operasyonu otomatik olarak başarısız kılmaz. NPC'nin ödülü azaltan ve ısıyı yükselten bir müdahale durumu üretmesini sağlar. Isı 0-100 aralığındadır, saatte 4 puan azalır ve oyuncuyu oyundan kovmak yerine risk profilini değiştirir.",
    )
    add_table(
        doc,
        ["Sonuç", "Temel ısı", "Kademe eklemesi", "Diğer etki"],
        [
            ("Başarı + gizli", "+4", "+2 x (kademe-1)", "Tam ödül"),
            ("Başarı + tespit", "+12", "+2 x (kademe-1)", "Ödül x0.72"),
            ("Başarısız + gizli", "+7", "+2 x (kademe-1)", "Kaynak harcanır, küçük analiz XP'si"),
            ("Başarısız + tespit", "+18", "+2 x (kademe-1)", "Kaynak harcanır, NPC alarmı uzar"),
        ],
        [2200, 1700, 2200, 3260],
        font_size=9.2,
        center_cols={1, 2},
    )
    add_figure(
        doc,
        "heat_detection_curve_v1_1.png",
        "Şekil 2. Isı arttıkça tespit riski yükselir; sistem sert enerji limiti kullanmaz.",
        "Dört hedef kademesi için 0 ile 100 ısı aralığında yükselen tespit olasılığı eğrileri.",
    )

    doc.add_heading("5.5 NPC olay müdahale durum makinesi", level=2)
    add_table(
        doc,
        ["Durum", "Trace bandı", "NPC davranışı", "Oyuncu karşılığı"],
        [
            ("Normal", "0-24", "Pasif telemetri toplar", "Hız avantajı"),
            ("Şüpheli", "25-49", "Ek doğrulama ve trafik sınırlama", "Maliyet veya rota artışı"),
            ("Alarm", "50-74", "Segment izolasyonu ve araç imza avı", "Geri çekilme kararı"),
            ("Containment", "75-99", "Kritik düğümleri kapatır, çıkışı daraltır", "Delili bırakıp temiz çıkış"),
            ("Kilitli", "100", "Operasyon sonlandırılır", "Kalıcı kayıp yok; rapor ve kısmi XP"),
        ],
        [1550, 1300, 3500, 3010],
        font_size=9.0,
        center_cols={1},
    )


def add_missions_and_ai(doc: Document) -> None:
    doc.add_heading("6. Yaşayan dünya, sözleşmeler, görevler ve NPC yapay zekâsı", level=1)
    doc.add_heading("6.1 İlk bölge ve NPC organizasyonları", level=2)
    add_body(
        doc,
        "Asteria Bölgesi; enerji, lojistik, belediye hizmetleri ve otonom üretimin birbirine bağlı olduğu kurgusal bir teknoloji koridorudur. NPC organizasyonları yalnız görev veren menüler değildir: kapasite yatırımı yapar, sözleşmelere teklif verir, pazar talebini değiştirir ve oyuncunun operasyon geçmişine tepki gösterir.",
    )
    add_table(
        doc,
        ["Organizasyon", "Stratejik amaç", "Savunma doktrini", "Oyuncuya sunduğu gerilim"],
        [
            ("Nexilune Industrial", "Otomasyon standardını tekelleştirmek", "Hızlı uyarlanan izleme ve analitik", "Kârlı fakat görünür sözleşmeler"),
            ("Asteria Civic Grid", "Hizmet sürekliliği ve kamu güveni", "Segmentasyon ve dayanıklılık", "Düşük sivil maliyetle çözüm"),
            ("Free Mesh", "Merkezi yapılardan bağımsız veri dolaşımı", "Düzensiz topoloji ve sahte sinyal", "Yüksek istihbarat, değişken risk"),
        ],
        [1800, 2700, 2600, 2260],
        font_size=8.9,
    )

    doc.add_heading("6.2 Dünya durumu ve sözleşme yarışı", level=2)
    add_list(
        doc,
        [
            "Dünya simülasyonu 6 saatlik mantıksal çevrimlerle arz, talep, hizmet kararlılığı ve NPC önceliklerini günceller.",
            "Ana hikâye sözleşmeleri teklif yarışına girmez; oyuncunun ilerlemesi rastgele NPC sonucuyla kilitlenmez.",
            "Pazar sözleşmelerinde oyuncu hazırlık puanı ile en iyi NPC teklifi karşılaştırılır; kayıp yalnız sınırlı başvuru teminatıdır.",
            "Kaçırılan fırsatlar yok olmaz; sözleşme panosunda en az 7 günlük arşiv hakkı bulunur.",
            "Ana merkez, araştırmalar ve kalıcı tesisler hiçbir dünya olayıyla geriye düşmez.",
        ],
    )
    add_formula(doc, "S_oyuncu = 0.30xhazırlık + 0.25xmaliyet_uyumu + 0.20xitibar + 0.15xistihbarat + 0.10xuzmanlaşma")
    add_formula(doc, "P_sözleşme = 0.10 + 0.80 x sigmoid((S_oyuncu - S_en_iyi_NPC)/9)")
    add_body(
        doc,
        "Olasılık %10-%90 arasında sınırlıdır. Teklif ekranı NPC puanını kesin sayı olarak değil açıklanabilir bir bant halinde gösterir. Kaybedilen teklif, sonraki hazırlık için rakip doktrin analizi üretir.",
    )

    doc.add_heading("6.3 MVP görev tipleri", level=2)
    add_table(
        doc,
        ["Görev", "Birincil statlar", "Oyuncu amacı", "Ana ödül"],
        [
            ("Delil Kurtarma", "Erişim, Gizlilik, Analiz", "Sentetik kanıt paketini bütünlüğü bozulmadan çıkar", "Delil + sermaye"),
            ("Hizmet Geri Kazanım", "Kontrol, Analiz", "Kilitli hizmet düğümlerini sırayla yeniden bağla", "Sermaye + bölgesel güven"),
            ("Atıf İzleme", "Keşif, Analiz, Gizlilik", "Düşman hücrenin imza zincirini tamamla", "Kampanya istihbaratı + araştırma"),
        ],
        [1900, 2200, 3400, 1860],
        font_size=9.0,
    )
    doc.add_heading("6.4 Hedef kademeleri", level=2)
    with BALANCE_PATH.open(encoding="utf-8") as handle:
        balance = json.load(handle)
    targets = balance["target_tiers"]
    add_table(
        doc,
        ["K", "Hedef", "Sert.", "İzleme", "Segm.", "Day.", "İşlem", "Enerji", "Sermaye", "Delil"],
        [
            (
                row["tier"], row["label"], row["hardening"], row["monitoring"], row["segmentation"],
                row["resilience"], row["compute_cost"], row["energy_cost"], row["base_capital"], row["base_evidence"],
            )
            for row in targets
        ],
        [450, 2400, 650, 650, 650, 650, 900, 800, 1000, 1210],
        font_size=7.8,
        center_cols={0, 2, 3, 4, 5, 6, 7, 8, 9},
    )
    add_body(
        doc,
        "Hedef istatistikleri oyuncu yüklemesini gördükten sonra değiştirilmez. Her hedef, kademe ortalaması + arketip değiştiricisi + seed'li [-5,+5] jitter ile oluşturulur ve seçim ekranına geldiği anda kalıcılaştırılır.",
    )
    add_formula(doc, "hedef_statı = clamp(kademe_min, kademe_max, kademe_ortalaması + arketip_modu + seed_jitter)")

    doc.add_heading("6.5 İlk kampanya: Asteria Enerji Krizi", level=2)
    add_list(
        doc,
        [
            "12 hedef varyantı: enerji dağıtımı, kimlik ağı, lojistik rölesi ve kamu hizmeti arketipleri.",
            "6 ana sözleşme ve 3 dünya olayı; her ana görev en az iki farklı stratejik hazırlıkla tamamlanabilir.",
            "3 çevresel değiştirici: bakım penceresi, talep sıçraması ve kısmi segment izolasyonu.",
            "1 kampanya finali: Nexilune otonom orkestratörü; mevcut kuralları birleştirir, yeni mini oyun eklemez.",
        ],
    )

    doc.add_heading("6.6 NPC hafızası ve doktrin adaptasyonu", level=2)
    add_body(
        doc,
        "Her organizasyon oyuncunun kullandığı yükleme etiketlerini ayrı ayrı izler. Aynı yaklaşımın tekrarı savunma puanını sınırlı biçimde yükseltir; kullanılmayan yaklaşımlar her kriz çevriminde %25 unutulur. Adaptasyon oyuncuya operasyondan önce gösterilir ve görev başladıktan sonra değişmez.",
    )
    add_formula(doc, "maruziyet_yeni = 0.75 x maruziyet_eski + bu_çevrimde_kullanım")
    add_formula(doc, "savunma_adaptasyonu = min(10, 2.5 x ln(1 + maruziyet))")
    add_table(
        doc,
        ["Maruziyet", "Savunma eklemesi", "Oyuncuya gösterim", "Amaç"],
        [
            (row["exposure"], f"+{row['defense_modifier']}", "Doktrin aşinalığı etiketi", "Aynı yüklemeyi sonsuza kadar baskın olmaktan çıkarmak")
            for row in balance["adaptation_table"]
        ],
        [1600, 2000, 2760, 3000],
        font_size=8.8,
        center_cols={0, 1},
    )

    doc.add_heading("6.7 Şeffaf zorluk yöneticisi", level=2)
    add_body(
        doc,
        "Gizli lastik bant zorluğu kullanılmaz. Sistem yalnızca oyuncuya sunulan hedef havuzunu ve açık yardım seçeneklerini düzenler. Aynı hedefin değerleri operasyon başladıktan sonra değişmez.",
    )
    add_table(
        doc,
        ["Sinyal", "Eşik", "Açık yardım", "Değişmeyen"],
        [
            ("Arka arkaya başarısızlık", "2", "+0.15 destek logiti ve ücretsiz ek önizleme", "Hedef statları ve ödül"),
            ("Yüksek ustalık", "Son 5 görevde >%85 başarı", "Bir üst kademe hedefi öner", "Mevcut görev olasılığı"),
            ("Düşük mini oyun skoru", "3 görev ortalaması <35", "40 sn erişilebilir mod ve otomatik nötr çözüm", "Ana başarı tabanı"),
            ("Uzun ara", "72 saat", "Bir defalık 2 istihbarat ve özet brifing", "Kalıcı ekonomi"),
        ],
        [2100, 1500, 3300, 2460],
        font_size=8.9,
        center_cols={1},
    )


def add_minigame(doc: Document) -> None:
    doc.add_heading("7. MVP mini oyunu: Node Routing", level=1)
    add_body(
        doc,
        "Node Routing, oyuncuya bir saldırı komutu yazdırmaz. Soyut bir ağ grafiğinde risk, gecikme ve delil değeri arasında rota seçtirir. Mini oyun ana stratejiyi destekler; atlanırsa nötr skor 50 uygulanır.",
    )
    add_table(
        doc,
        ["Parametre", "Standart", "Erişilebilir mod"],
        [
            ("Süre", "25 saniye", "40 saniye veya süre kapalı alıştırma"),
            ("Graf", "5 sütun, sütun başına 3-4 düğüm, 22-28 kenar", "Aynı; daha güçlü odak ve rota önizlemesi"),
            ("Kontrol", "Fare + klavye", "Tam klavye, tek tuş geri alma"),
            ("Atlama", "Skor 50", "Skor 50; ödül cezası yok"),
            ("Görsel", "Renk + ikon + metin etiketi", "Azaltılmış hareket, yükseltilmiş kontrast"),
        ],
        [2250, 3500, 3610],
        font_size=9.0,
    )
    add_formula(doc, "skor = 100 x [0.55x(1-ortalama_risk) + 0.25x(1-ortalama_gecikme) + 0.20xpaket_oranı]")
    add_formula(doc, "beceri_bonus = clamp(-0.35, +0.40, (skor-50)/125)")
    add_formula(doc, "skor_gürültüsü = 6 - 0.12 x skor")
    add_body(
        doc,
        "Skor 100 başarı logitine en fazla +0.40 ekler ve gürültüyü 6 puan azaltır. Skor 0 en fazla -0.35 ve +6 gürültü getirir. Eşit güçte bir hedefte bu yaklaşık 5-9 yüzde puanlık başarı farkıdır; stratejik hazırlığı geçersiz kılmaz.",
    )

    doc.add_heading("7.1 Üretim kabul ölçütleri", level=2)
    add_list(
        doc,
        [
            "Her üretilen grafın başlangıçtan hedefe en az bir geçerli yolu vardır.",
            "En iyi yol tek bir sabit algoritmaya indirgenmez; risk, süre ve paket değeri gerçek ödünleşim yaratır.",
            "Aynı seed aynı grafı üretir; tekrar ve hata raporu yeniden oynatılabilir.",
            "60 FPS hedeflenir; metin her karede yeniden oluşturulmaz ve görünmeyen düğümler çizilmez. [R12]",
            "Renk tek bilgi kanalı değildir; klavye odağı görünür ve hareket azaltma tercihi desteklenir. [R16]",
        ],
    )


def add_economy(doc: Document) -> None:
    doc.add_heading("8. Ekonomi, ilerleme ve bütün matematik", level=1)
    with BALANCE_PATH.open(encoding="utf-8") as handle:
        balance = json.load(handle)

    doc.add_heading("8.1 Kaynak mimarisi", level=2)
    add_table(
        doc,
        ["Ana kaynak", "Başlıca kaynak", "Başlıca tüketim", "Tasarım görevi"],
        [
            ("Enerji", "Mikro Şebeke", "Tesis öncelikleri ve operasyon", "Aynı anda her sistemi tam güçte çalıştırmama kararı"),
            ("İşlem Gücü", "Veri Merkezi", "Keşif, araştırma ve operasyon", "Hazırlık ile fırsat maliyetini bağlamak"),
            ("Bileşen", "Robotik Atölye", "Tesis ve araç modülü", "İnşa zinciri ve uzmanlaşma"),
            ("Sermaye", "Sözleşme, pazar ve delil satışı", "Yükseltme, personel ve teklif teminatı", "Aktif karar ile kalıcı yatırım arasında köprü"),
            ("Uzmanlık", "Araştırma Laboratuvarı ve ilk tamamlama", "Araştırma düğümleri", "Ham güce değil yeni seçeneğe ilerlemek"),
        ],
        [1800, 2100, 2500, 2960],
        font_size=8.9,
    )
    add_table(
        doc,
        ["Operasyon durumu", "Tür", "Nasıl değişir?", "Neden ana kaynak değil?"],
        [
            ("Bant kapasitesi", "Kapasite", "Veri Merkezi, SOC ve modüllerle artar", "Yükleme sınırıdır; harcanan para değildir"),
            ("İstihbarat", "Hazırlık", "Keşif, başarısızlık analizi ve sözleşme", "Belirsizliği azaltan bağlamsal avantajdır"),
            ("Delil", "Görev çıktısı", "Operasyon ve adli rapor", "Satış, araştırma veya ilişki için seçim nesnesidir"),
            ("Isı", "Risk", "Tespit ile artar, zaman/SOC ile azalır", "Oyuncuyu durduran enerji duvarı değildir"),
            ("Etki", "Kampanya puanı", "Kriz çözümü ve sözleşme sonucu", "Zafer yönünü ölçer; satın alma gücü değildir"),
        ],
        [1900, 1500, 3000, 2960],
        font_size=8.8,
    )
    add_formula(doc, "bant_kapasitesi = 20 + 6xVeri_Merkezi_sv + 4xSOC_sv + modül_bonusları")
    add_callout(
        doc,
        "Birim kararı",
        "Arayüz atmosfer için MW, PFLOPS veya Gbps gösterebilir; sunucu ekonomisi normalize tam sayılarla çalışır. Fiziksel olarak farklı birimler sahte kesinlikle birbirine çevrilmez.",
        "info",
    )

    doc.add_heading("8.2 Tesis ağı ve yükseltme eğrisi", level=2)
    add_formula(doc, "çıktı(L) = taban_çıktı x 1.24^(L-1)")
    add_formula(doc, "sermaye_maliyeti(L) = taban_sermaye x 1.55^(L-1)")
    add_formula(doc, "bileşen_maliyeti(L) = taban_bileşen x 1.48^(L-1)")
    add_formula(doc, "süre(L) = min(360, taban_süre x 1.50^(L-1)) dakika")
    facility_rows = balance["facility_table"]
    facility_names = [
        "Mikro Şebeke",
        "Veri Merkezi",
        "Robotik Atölye",
        "Araştırma Laboratuvarı",
        "Güvenlik Operasyon Merkezi",
    ]
    facility_table_rows = []
    for name in facility_names:
        l1 = next(row for row in facility_rows if row["facility"] == name and row["level"] == 1)
        l5 = next(row for row in facility_rows if row["facility"] == name and row["level"] == 5)
        facility_table_rows.append(
            (
                name,
                l1["output_kind"],
                l1["output_per_hour"],
                l5["output_per_hour"],
                l5["upgrade_capital_cost"],
                l5["upgrade_component_cost"],
                l5["upgrade_time_minutes"],
            )
        )
    add_table(
        doc,
        ["Tesis", "Çıktı", "L1/s", "L5/s", "L5 serm.", "L5 bil.", "L5 süre"],
        facility_table_rows,
        [1900, 1500, 850, 850, 1250, 1200, 1810],
        font_size=8.2,
        center_cols={2, 3, 4, 5, 6},
    )
    add_body(
        doc,
        "Tesis seviyesi 12 ile sınırlıdır. İleri gelişim aynı binayı sonsuza kadar büyütmek yerine modül, personel doktrini ve üretim önceliği açar. İnşa ve araştırma kuyrukları ücretsizdir; iptal edilen proje başlamadıysa %100, başladıysa harcanmamış pay üzerinden en az %90 iade verir.",
    )
    add_figure(
        doc,
        "facility_curve_v1_1.png",
        "Şekil 3. Veri Merkezi üretimi ile sermaye maliyeti farklı eğrilerde büyür; süre altı saatte sınırlandırılır.",
        "Birden on ikiye Veri Merkezi seviyelerinde normalize üretim hızı ve yükseltme sermaye maliyeti.",
    )

    doc.add_heading("8.3 Enerji önceliği ve çevrimdışı üretim", level=2)
    add_formula(doc, "güç_oranı = min(1, ayrılan_enerji / gereken_enerji)")
    add_formula(doc, "etkin_çıktı = nominal_çıktı x güç_oranı x personel_oranı x durum_oranı")
    add_body(
        doc,
        "Enerji yetmediğinde tesisler rastgele kapanmaz. Oyuncu öncelik sırası belirler; kritik tesisler önce beslenir ve diğerleri kısmi verimde çalışır. Bu, enerji krizini yalnız kırmızı eksi sayı değil gerçek stratejik karar yapar.",
    )
    add_formula(doc, "delta_saat = min((şimdi - son_hesap)/3600, depolama_saati)")
    add_formula(doc, "yeni_bakiye = min(kapasite, eski_bakiye + etkin_saatlik_üretim x delta_saat)")
    add_formula(doc, "depolama_saati(L) = min(36, 24 + 1.5x(L-1))")
    add_body(
        doc,
        "Üretim okuma veya işlem anında son zaman damgasından hesaplanır; oyuncu başına sürekli cron gerekmez. Başlangıç depolaması 24 saattir ve 36 saate çıkar. Bu karar eski sekiz saatlik tavanı kaldırır; günde birden fazla giriş baskısı oluşturmaz.",
    )

    doc.add_heading("8.4 Operasyon maliyeti, ödül ve başarısızlık değeri", level=2)
    add_table(
        doc,
        ["Yükleme", "İşlem modu", "Enerji modu", "Bant yükü", "Araç / görünürlük"],
        [
            ("Sessiz İnceleme", "0.90", "0.90", "Düşük", "+3 araç; -4 gürültü, +8 örtü"),
            ("Dengeli Ajan", "1.00", "1.00", "Orta", "+6 araç; nötr gürültü, +4 örtü"),
            ("Hızlı Köprü", "1.20", "1.15", "Yüksek", "+10 araç; +8 gürültü, -2 örtü"),
        ],
        [2100, 1400, 1400, 1500, 2960],
        font_size=8.8,
        center_cols={1, 2, 3},
    )
    add_formula(doc, "işlem_maliyeti = kademe_işlem x görev_modu x yükleme_modu")
    add_formula(doc, "enerji_maliyeti = kademe_enerji x görev_modu x yükleme_modu")
    add_formula(doc, "kalite = clamp(0.65, 1.15, 0.80 + 0.20xP_başarı + 0.003x(skor-50))")
    add_formula(doc, "sermaye_ödülü = taban_sermaye x kalite x gizlilik x tekrar x pazar x U(0.92,1.08)")
    add_formula(doc, "gizlilik = 1.00 (tespit yok), 0.72 (tespit var)")
    add_formula(doc, "tekrar = max(0.55, 1 - 0.15 x son_3_görevde_aynı_arketip_sayısı)")
    add_formula(doc, "başarısız_delil = taban_delil x 0.25 x [0.80 tespit varsa, aksi halde 1.00]")
    add_body(
        doc,
        "Başarısızlık boş oturum değildir: en az %25 taban delil analiz verisine dönüşür. Erişim öncesi geri çekilmede maliyetin %70'i, erişim sonrasında %30'u iade edilir; hedef aşamasından sonra iade yoktur. Kalıcı tesis veya araştırma kaybı oluşmaz.",
    )

    doc.add_heading("8.5 NPC pazarı ve fiyat istikrarı", level=2)
    add_formula(doc, "pazar_t = clamp(0.85, 1.15, 1 + 0.72x(pazar_(t-1)-1) + şok_t)")
    add_formula(doc, "fiyat = taban_fiyat x pazar_endeksi x kıtlık x sözleşme_kalitesi")
    add_body(
        doc,
        "MVP pazarı oyuncular arası değildir. NPC arz-talebi ve hikâye şokları fiyatları değiştirir; 0.85-1.15 koridoru kontrolsüz enflasyonu engeller. Delili satmak kısa vadeli sermaye, araştırmaya ayırmak uzun vadeli seçenek, bir kuruma vermek ilişki ve etki üretir.",
    )
    add_figure(
        doc,
        "market_index_curve_v1_1.png",
        "Şekil 4. Sabit tohumlu 120 günlük NPC pazarı iki senaryo şokuna rağmen denge koridoruna geri döner.",
        "Yüz yirmi simülasyon gününde 85 ile 115 arasında dalgalanan ve 100 çevresine dönen NPC pazar endeksi.",
    )

    doc.add_heading("8.6 Aktif-casual güç farkı", level=2)
    add_formula(doc, "aktif_ekonomi_bonusu(n) = min(0.20, 0.08xln(1+n))")
    add_table(
        doc,
        ["Profil", "Operasyon/gün", "Ekonomik bonus", "Günlük ilerleme endeksi"],
        [
            (row["profile"], row["operations_per_day"], f"%{row['economic_bonus']*100:.1f}", row["daily_progress_index"])
            for row in balance["activity_profiles"]
        ],
        [2200, 2100, 2300, 2760],
        font_size=9.0,
        center_cols={1, 2, 3},
    )
    add_body(
        doc,
        "Günde 2 operasyon ile 10 operasyon profilleri arasındaki ekonomik hız farkı modelde yaklaşık %9.5'tir. Uzun oynayan kişi daha fazla ustalık, hikâye, bilgi ve kozmetik ilerleme kazanır; ham ekonomi avantajı %20 tavanını aşmaz.",
    )
    add_figure(
        doc,
        "activity_gap_curve_v1_1.png",
        "Şekil 5. Aktif katkı ilk operasyonlarda değerlidir, sonra %20 ekonomik tavanına yaklaşır.",
        "Günlük operasyon sayısı arttıkça azalan marjinal faydayla yüzde yirmide sınırlanan ekonomik hız bonusu.",
    )

    doc.add_heading("8.7 Uzmanlık ve seviye eğrisi", level=2)
    add_formula(doc, "XP_sonraki(L) = round(60 x L^1.35)")
    add_formula(doc, "XP_görev = round(40 x kademe^1.35 x sonuç x ilk_tamamlama x zorluk)")
    add_body(
        doc,
        "Sonuç çarpanı başarıda 1.0, başarısızlıkta 0.25; ilk tamamlama 1.20; zorluk 0.80-1.25 aralığındadır. Seviye başarıya gizli bonus vermez; araştırma görünürlüğü ve yatay seçenek açar.",
    )
    xp = balance["xp_table"][:12]
    add_table(
        doc,
        ["Seviye", "Sonraki seviye XP", "Kümülatif XP"],
        [(row["level"], row["xp_to_next"], row["cumulative_xp"]) for row in xp],
        [1800, 3600, 3960],
        font_size=9.0,
        center_cols={0, 1, 2},
    )
    add_formula(doc, "araştırma_uzmanlık_maliyeti(rütbe) = round(8 x 1.45^(rütbe-1)); araştırma_delil_maliyeti = round(6 x 1.50^(rütbe-1))")
    add_formula(doc, "ham_stat = taban_stat + araştırma_bonusları + yükleme_bonusları; ardından effective_stat uygulanır")
    add_body(
        doc,
        "Bir araştırma düğümü çoğunlukla yeni araç, otomasyon kuralı veya yan avantaj/dezavantaj açar. Düz istatistik düğümleri birincil stata en fazla +3 verir ve aynı dalda ardışık istif sınırlıdır.",
    )

    doc.add_heading("8.8 Monte Carlo denge sonucu", level=2)
    matched = balance["matched_simulation"]
    add_table(
        doc,
        ["Profil", "Kad.", "Saldırı", "Savunma", "Başarı", "Tespit", "Gizli başarı", "Bek. serm.", "Bek. ısı"],
        [
            (
                row["player"], row["target_tier"], row["attack_rating"], row["defense_rating"],
                f"%{row['success_rate']*100:.1f}", f"%{row['detection_rate']*100:.1f}",
                f"%{row['success_and_undetected_rate']*100:.1f}", row["expected_capital_per_attempt"], row["expected_heat_gain"],
            )
            for row in matched
        ],
        [1200, 520, 920, 920, 920, 850, 1300, 1220, 1510],
        font_size=7.8,
        center_cols={1, 2, 3, 4, 5, 6, 7, 8},
    )
    add_body(
        doc,
        "Simülasyon sabit seed 20260809 ile eş-kademe çift başına 100.000 koşu kullanır. Kademe 1-3 yaklaşık %77-79 başarı; kademe 4-5 artan tespit ve risk üretir. Bir üst kademeye erken giriş %48-61 bandındadır. Bunlar alfa öncesi başlangıç hedefidir, oyuncu testi yerine geçmez.",
    )

    doc.add_heading("8.9 Etki ve kampanya zaferi", level=2)
    add_formula(doc, "etki_kazanımı = taban_etki x kalite x kriz_katkısı x ilk_çözüm")
    add_body(
        doc,
        "Kampanya tek bir düşmanı yok ederek bitmez. Oyuncu dört başarı ekseninden üçünü hedefe ulaştırır: Bölgesel Güven, Teknoloji Standardı, Ekonomik Etki ve İstihbarat Üstünlüğü. Bu yapı farklı doktrinlerin aynı kampanyayı farklı biçimde tamamlamasını sağlar.",
    )
    add_callout(doc, "Matematik denetimi", "; ".join(balance["audit_checks"]), "positive")


def add_retention(doc: Document) -> None:
    doc.add_heading("9. Uzun süreli bağlılık ve etik içerik ritmi", level=1)
    add_body(
        doc,
        "Oyun bağlılığı, SDT'nin özerklik, yeterlik ve bağ kurma ihtiyaçlarına dayandırılır. Kalıcı dünya tek başına yeterli değildir; oyuncu dünyayı okuyabilmeli, farklı doktrinler deneyebilmeli ve kararının sonuçlarını sonraki oturumda görebilmelidir. Kontrol ustalığı gereklidir ancak tek başına yeterli değildir. [R5-R7, R17-R18]",
    )
    add_table(
        doc,
        ["İhtiyaç", "Tasarım karşılığı", "Kaçınılacak desen"],
        [
            ("Özerklik", "Sözleşme, üretim önceliği, rota, yükleme ve delil kullanımı", "Tek doğru yapı ve zorunlu günlük görev"),
            ("Yeterlik", "Açıklanabilir rapor, ustalık skorları, tekrar oynatma", "Gizli zorluk değişimi ve belirsiz kayıp"),
            ("Bağ kurma", "NODE-7 ekibi, NPC ilişkileri, kurtarılan hizmetler ve bölge tarihi", "Çevrimdışı PvP baskısı veya ücretli ayrıcalık"),
        ],
        [1700, 4700, 2960],
        font_size=9.2,
    )

    doc.add_heading("9.1 İçerik katmanları", level=2)
    add_list(
        doc,
        [
            "Anlık: pazar sinyali, hedef grafı, araç eşleşmesi, Node Routing ve çıkış kararı.",
            "Oturumluk: 1-3 sözleşme, üretim önceliği, üs yatırımı ve dünya tepkisi.",
            "Krizlik: 3-6 bağlantılı görev; enerji arzı, NPC doktrini ve bölgesel güven üzerinde sonuç.",
            "Uzun vadeli: yeni bölgeler, ekip karakterleri, alternatif tesis/araştırma yapıları ve kozmetik merkez temaları.",
        ],
    )
    add_callout(
        doc,
        "Etik sınır",
        "Günlük seri kaybı, çevrimdışı yağma, loot box, yapay kıtlık, ücretli saldırı bildirimi, kaybolan sezon gücü ve oyuncuyu çıkmamaya zorlayan sayaç kullanılmaz. Sözleşmeler en az 7 gün arşivlenir; bildirimlerin temel işlevleri ücretsiz olur.",
        "caution",
    )

    doc.add_heading("9.2 Prototip deney hedefleri", level=2)
    add_table(
        doc,
        ["Ölçüm", "Hedef", "Karar"],
        [
            ("İlk görev başlatma", ">=%80", "Altındaysa açılış ve bilgi yoğunluğu sadeleştirilir"),
            ("İlk görev tamamlama", ">=%70", "Altındaysa yönlendirme ve zorluk bandı gözden geçirilir"),
            ("Sonucu açıklayabilme", ">=%80 doğru neden", "Altındaysa rapor dili/formülü değişir"),
            ("İkinci görevi gönüllü başlatma", ">=%60", "Altındaysa çekirdek döngü içerikten önce yeniden tasarlanır"),
            ("Strateji-operasyon bağını açıklama", ">=%75", "Altındaysa dünya sonucu ve kaynak geri bildirimi sadeleştirilir"),
            ("Algılanan adalet", ">=4/5", "Altındaysa RNG, telemetri ve başarı tahmini incelenir"),
            ("Sağlıklı oturum", "Medyan 20-35 dk", "Daha uzun süre otomatik olarak başarı sayılmaz"),
        ],
        [2650, 1700, 5010],
        font_size=9.0,
        center_cols={1},
    )
    add_body(doc, "Bu eşikler pazar benchmarkı değil, dikey kesit için ürünün kendi go/no-go hedefleridir.")


def add_ui_ux(doc: Document) -> None:
    doc.add_heading("10. UI/UX ve erişilebilirlik spesifikasyonu", level=1)
    doc.add_heading("10.1 Ana ekran bilgi mimarisi", level=2)
    add_table(
        doc,
        ["Bölge", "İçerik", "Öncelik"],
        [
            ("Üst durum şeridi", "Enerji, İşlem, Bileşen, Sermaye, Uzmanlık ve Isı", "Birincil kaynaklar ile risk görsel olarak ayrılır"),
            ("Sol modül navigasyonu", "Dünya, Merkez, Sözleşmeler, Operasyon, Araştırma", "Klavye ile erişilir; en fazla 5 ana modül"),
            ("Merkez çalışma alanı", "2D bölge ağı, tesis görünümü veya operasyon grafı", "PixiJS; kritik bilgi DOM panellerinde yinelenir"),
            ("Sağ bağlam paneli", "Pazar/kriz etkisi, hedef, başarı bandı, yükleme ve olay akışı", "Seçime göre değişir; sabit 360-420 px"),
            ("Alt eylem alanı", "Başlat, geri çekil, rapor, yardım", "Tek birincil eylem; risk öncesi açık onay"),
        ],
        [2000, 4500, 2860],
        font_size=9.1,
    )
    add_body(
        doc,
        "İlk sürüm 3D küre kullanmaz. Bölge, tesis ve ağ topolojileri okunabilir 2D/2.5D görünümlerdir. PixiJS yalnız dinamik harita ve mini oyun için kullanılır; metin, form, pazar ve kritik etkileşimler erişilebilir DOM bileşenlerinde kalır.",
    )

    doc.add_heading("10.2 Görsel dil", level=2)
    add_table(
        doc,
        ["Token", "Değer", "Kullanım"],
        [
            ("Arka plan", "#0A0D14", "Ana kabuk"),
            ("Güvenli", "#00FF66", "Başarı; yalnız renkle değil ikon/metinle"),
            ("Bilgi", "#00E5FF", "Bağlantı ve seçili öğe"),
            ("Uyarı", "#FFB800", "Orta risk ve devam eden eylem"),
            ("Kritik", "#FF0055", "Yüksek risk, hata ve containment"),
            ("Metin", "#E6EDF3 / #9BA7B4", "Birincil / ikincil"),
            ("Font", "Inter + JetBrains Mono", "Arayüz / kısa veri"),
        ],
        [1750, 1900, 5710],
        font_size=9.2,
    )
    add_list(
        doc,
        [
            "WCAG 2.2 AA hedeflenir; klavye odağı görünür ve odaklanan öğe örtülmez. [R16]",
            "Renk tek anlam taşımaz; kritik durumlarda ikon, kısa etiket ve sayı birlikte kullanılır.",
            "prefers-reduced-motion etkinse büyük ölçekleme, parallax ve sürekli titreşim kaldırılır.",
            "Mini oyun sürükleme gerektiren tek çözüm sunmaz; tıklama ve klavye alternatifi vardır.",
            "SFX, ekran titremesi, scanline/noise ve otomatik kayan log ayrı ayrı kapatılabilir.",
        ],
    )


def add_technical_architecture(doc: Document) -> None:
    doc.add_heading("11. Teknik mimari", level=1)
    add_callout(
        doc,
        "Mimari kararı",
        "MVP mikroservis değil modüler monolit olacaktır: tek API dağıtımı, ayrı worker süreci, PostgreSQL sistem kaydı ve Redis geçici/iş kuyruğu katmanı. Modül sınırları daha sonra ayrılmaya uygun tutulur.",
        "positive",
    )
    add_table(
        doc,
        ["Katman", "Öneri", "Sorumluluk"],
        [
            ("Web istemci", "React + TypeScript", "Erişilebilir DOM arayüzü, UI durum yönetimi, form ve raporlar"),
            ("Görsel motor", "PixiJS 8 / WebGL", "Ağ haritası ve Node Routing; üretimde WebGL, WebGPU deneysel kalır [R12]"),
            ("API", "NestJS modüler monolit", "Kimlik, profil, ekonomi, görev ve sonuç API'leri"),
            ("Worker", "NestJS + BullMQ", "Yükseltme/operasyon bitişi, bildirim ve uzlaşma işleri [R13]"),
            ("Sistem kaydı", "PostgreSQL 18", "Bakiyeler, operasyonlar, olay günlüğü ve içerik sürümleri; transaction sınırları [R14]"),
            ("Geçici katman", "Redis 8.x", "Cache, rate limit, kısa ömürlü oturum, BullMQ; bakiye otoritesi değildir [R15]"),
            ("Gerçek zaman", "WebSocket/SSE", "Durum bildirimleri; istemci yeniden bağlandığında cursor ile eksik olayları alır"),
            ("Dağıtım", "Docker + reverse proxy", "Yerel geliştirme ve tek bölge beta; gözlemlenebilirlik dahil"),
        ],
        [1550, 2300, 5510],
        font_size=8.8,
    )

    doc.add_heading("11.1 Sunucu otoritesi ve deterministik çözüm", level=2)
    add_list(
        doc,
        [
            "İstemci yalnız niyet gönderir; bakiye, başarı, ödül, seed ve zaman sunucuda hesaplanır.",
            "Her operasyon başlatılırken içerik sürümü, formül sürümü ve kriptografik olarak güçlü seed kaydedilir.",
            "Node Routing grafı seed ile yeniden üretilebilir; istemcinin skor girdisi sunucu kurallarıyla doğrulanır.",
            "Her harcama idempotency key taşır; aynı istek iki kez bakiye düşüremez.",
            "PostgreSQL transaction, bakiye değişimi ile ledger olayını atomik kaydeder.",
            "Queue kaçırılırsa finish_at alanına göre API/uzlaşma worker'ı işlemi tamamlar; Redis tek doğruluk kaynağı değildir.",
        ],
    )

    doc.add_heading("11.2 Ana veri modeli", level=2)
    add_table(
        doc,
        ["Tablo", "Kritik alanlar", "Kural"],
        [
            ("users", "id, handle, auth_subject, created_at", "PII minimum; oyun profili kimlik sağlayıcıdan ayrılır"),
            ("profiles", "user_id, level, xp, heat, region, doctrine", "Tek oyuncu ilerleme özeti"),
            ("resource_ledger", "id, user_id, kind, delta, balance_after, reason, idem_key", "Append-only; her değişimin nedeni vardır"),
            ("production_state", "user_id, energy_priority, rates, capacities, last_accrual_at", "24-36 saatlik lazy accrual"),
            ("facilities", "user_id, type, level, power_priority, finish_at, version", "Ücretsiz inşa ve araştırma kuyrukları"),
            ("world_state", "region, cycle, demand, stability, active_crises, version", "Altı saatlik deterministik dünya çevrimi"),
            ("factions", "id, doctrine, capacity, priorities, player_relation", "NPC strateji ve ilişki durumu"),
            ("market_ticks", "cycle, commodity, index, supply, demand, shock", "0.85-1.15 endeks koridoru"),
            ("contracts", "id, issuer, objective, bid_band, impact, archive_until", "Hikâye ve pazar sözleşmeleri ayrılır"),
            ("targets", "id, template_id, tier, seed, stats, expires_at", "Seçimden önce sabitlenir"),
            ("operations", "id, user_id, target_id, state, seed, formula_version", "Durum makinesi ve tekrar oynatma"),
            ("operation_events", "operation_id, seq, type, payload, created_at", "Sıralı olay kaydı"),
            ("loadouts", "user_id, tool_ids, revision", "Sunucu doğrulamalı araç seti"),
            ("content_versions", "version, checksum, published_at", "Denge ve hedef içeriği sürümlenir"),
        ],
        [1700, 4300, 3360],
        font_size=8.5,
    )

    doc.add_heading("11.3 Modül sınırları", level=2)
    add_table(
        doc,
        ["Modül", "Komutlar", "Yayınladığı olaylar"],
        [
            ("Identity", "register, login, logout", "user.created, session.revoked"),
            ("Economy", "accrue, allocatePower, spend, grant, sellEvidence", "resource.changed, production.rebalanced"),
            ("Facility", "startUpgrade, claimUpgrade", "facility.started/completed"),
            ("World", "advanceCycle, applyCrisis, updateFaction", "world.cycle, crisis.changed, faction.acted"),
            ("Market", "quote, settleNpcTrade", "market.tick, market.trade"),
            ("Contract", "list, bid, accept, archive", "contract.offered/awarded/expired"),
            ("Mission", "listTargets, recon, startOperation", "operation.started"),
            ("Simulation", "submitDecision, submitMinigame, resolve", "operation.phase/resolved"),
            ("Content", "publishVersion, activateVersion", "content.activated"),
            ("Telemetry", "ingestBatch", "analytics.accepted"),
        ],
        [1700, 3600, 4060],
        font_size=8.8,
    )


def add_security_and_ethics(doc: Document) -> None:
    doc.add_heading("12. Ürün güvenliği, kötüye kullanım ve etik", level=1)
    add_body(
        doc,
        "Siber güvenlik temalı bir ürünün kendi güvenlik seviyesi özellikle önemlidir. MVP, OWASP ASVS 5.0 Level 2'yi geliştirme ve kabul kontrol listesi olarak, NIST SSDF 1.1'i süreç çerçevesi olarak kullanır. [R8-R10]",
    )
    add_list(
        doc,
        [
            "Kimlik doğrulama güvenli, HttpOnly ve SameSite oturum çerezleriyle; hassas eylemler yeniden yetkilendirmeyle korunur.",
            "Her REST ve WebSocket mesajı şema ile doğrulanır; yetki bağlantı seviyesinde değil mesaj seviyesinde kontrol edilir.",
            "WebSocket Origin allowlist, hız sınırı, mesaj boyutu sınırı, nonce/idempotency ve oturum süresi kontrolü uygular. [R9]",
            "İstemciden gelen süre, skor, bakiye ve seed güvenilir kabul edilmez.",
            "Yönetim paneli ayrı rol, ayrı audit log ve mümkünse ayrı ağ politikası kullanır.",
            "Bağımlılık kilidi, SBOM, secret scanning, SAST ve düzenli yedek geri dönüş testi CI/CD kabul kapısıdır.",
            "Telemetri, sohbet veya kullanıcı içeriğinde gerçek saldırı hedefi/komut paylaşımını teşvik edecek alanlar MVP'de bulunmaz.",
        ],
    )
    add_callout(
        doc,
        "Kırmızı çizgi",
        "Oyunun tema uğruna gerçek saldırı kapasitesi üretmesine izin verilmez. İçerik tasarımı, güvenlik incelemesi ve yayın öncesi kötüye kullanım testi ayrı kabul kapısıdır.",
        "risk",
    )


def add_telemetry_and_balance(doc: Document) -> None:
    doc.add_heading("13. Telemetri, dengeleme ve test bilimi", level=1)
    doc.add_heading("13.1 Zorunlu olaylar", level=2)
    add_table(
        doc,
        ["Olay", "Temel alanlar", "Soru"],
        [
            ("tutorial_step", "step, duration, skipped, help_opened", "Oyuncu nerede takılıyor?"),
            ("world_viewed", "cycle, crisis, market_signal, duration", "Oyuncu dünya durumunu okuyabiliyor mu?"),
            ("production_rebalanced", "before, after, shortage, expected_output", "Enerji önceliği anlamlı karar mı?"),
            ("contract_selected", "issuer, objective, bid_band, predicted_impact", "Sözleşme seçiminde ne belirleyici?"),
            ("target_viewed", "tier, uncertainty, predicted_p", "Risk bilgisi okunuyor mu?"),
            ("loadout_changed", "tools, cost, noise, predicted_p", "Araç seçimi anlamlı mı?"),
            ("operation_started", "target, seed, formula_version, balances", "Başlangıç durumu yeniden üretilebilir mi?"),
            ("phase_decision", "phase, choice, trace_before/after", "Hangi kararlar dominant?"),
            ("minigame_result", "score, duration, path_risk, input_mode", "Beceri ve erişilebilirlik dengesi nasıl?"),
            ("operation_resolved", "success, detected, reward, heat, world_delta, explanation", "Model ve dünya sonucu anlaşılır mı?"),
            ("world_effect_applied", "market_delta, relation_delta, stability_delta, influence", "Taktik sonuç stratejiye bağlanıyor mu?"),
            ("session_end", "duration, operations, voluntary_exit", "Sağlıklı oturum oluşuyor mu?"),
        ],
        [1800, 4700, 2860],
        font_size=8.7,
    )
    add_body(
        doc,
        "Kişisel veri yerine olay bağlamı toplanır. Ham rota tıklamaları kısa saklanır; uzun vadeli analizde özet istatistikler kullanılır. Her telemetri olayı schema_version taşır.",
    )

    doc.add_heading("13.2 Dengeleme yöntemi", level=2)
    add_list(
        doc,
        [
            "Önce deterministik birim testleri: formül sınırları, monotonluk, tam sayı ve overflow kontrolleri.",
            "Ardından Monte Carlo: kademe/profil matrisi, ısı duyarlılığı, ödül varyansı ve en iyi strateji araması.",
            "Sonra bot politika testi: sessiz, hızlı, dengeli, ekonomik ve istihbarat odaklı ajanların beklenen değeri.",
            "Ekonomi stres testi: 120 günlük NPC pazarı, enerji kıtlığı ve günde 0-30 operasyon profili.",
            "Son olarak insan testi: açıklanabilirlik, adalet, karar çeşitliliği ve ikinci operasyon isteği.",
            "Bir değişken bir sürümde tek başına değiştirilir; formula_version ile eski operasyon kayıtları korunur.",
        ],
        ordered=True,
    )
    add_table(
        doc,
        ["Alarm", "Eşik", "Muhtemel müdahale"],
        [
            ("Dominant yükleme", "Tek yükleme seçimlerin >%55'i", "Karşı hedef/modifier veya araç gürültü maliyeti"),
            ("Aşırı şans", "Açıklanan p ile sonuç kalibrasyon hatası >3 puan", "PRNG, örneklem ve formül denetimi"),
            ("Ekonomi şişmesi", "Medyan sermaye 7 günde hedefin >%25 üstü", "Sinks ve tekrar çarpanı; ödülü körlemesine kısmama"),
            ("Pazar sapması", "Endeks 0.85/1.15 sınırında 3+ çevrim", "Ortalama dönüş veya şok genliğini ayarla"),
            ("Aktif güç uçurumu", "2 ve 10 operasyon/gün profili farkı >%20", "Aktif bonus eğrisini daralt"),
            ("Isı duvarı", "Isı nedeniyle gönüllü çıkışların >%20'si", "Decay veya katkı katsayısını düşür"),
            ("Mini oyun zorunluluğu", "Skip kullananların başarı farkı >10 puan", "Beceri bonusunu daralt"),
        ],
        [2100, 3400, 3860],
        font_size=8.9,
    )


def add_roadmap(doc: Document) -> None:
    doc.add_heading("14. Üretim yol haritası ve ekip planı", level=1)
    add_table(
        doc,
        ["Faz", "Süre", "Kapsam", "Çıkış kapısı"],
        [
            ("0. Ön üretim", "3 hafta", "Birleşik döngü, ekonomi modeli, UX akışı, sanat yönü, içerik güvenliği", "Kağıt prototipte strateji-operasyon bağı"),
            ("1. Teknik temel", "4 hafta", "Monorepo, auth, DB, ledger, lazy accrual, içerik sürümü", "Çifte harcama/idempotency testleri geçer"),
            ("2. Strateji prototipi", "5 hafta", "Beş kaynak, enerji önceliği, NPC pazarı, sözleşme ve dünya çevrimi", "Dünya kararı operasyon hazırlığını değiştirir"),
            ("3. Taktik dikey kesit", "6 hafta", "Keşif, yükleme, Node Routing, resolver ve adli rapor", "20 dakikalık uçtan uca oynanış"),
            ("4. Birleşik MVP", "6 hafta", "3 NPC, 3 görev, 12 hedef, kampanya, araştırma ve dünya sonuçları", "Operasyon pazarı/ilişkiyi güvenilir değiştirir"),
            ("5. Sertleştirme", "4 hafta", "Erişilebilirlik, güvenlik, performans, telemetri", "ASVS L2 kapsamı ve WCAG kritik akışlar"),
            ("6. Kapalı alfa", "3 hafta", "50-100 oyuncu, ekonomi, denge ve öğretici iterasyonu", "Go/no-go deney hedefleri"),
        ],
        [1600, 1150, 4000, 2610],
        font_size=8.8,
        center_cols={1},
    )
    add_body(
        doc,
        "Önerilen takvim 3 kişilik çekirdek ekipte yaklaşık 31 haftadır: 1 full-stack/backend-simülasyon, 1 frontend/gameplay, 1 ürün/UI-içerik; QA, ses ve güvenlik incelemesi yarı zamanlı desteklenir. Tek deneyimli geliştirici + AI desteği için gerçekçi tam zamanlı aralık 48-64 haftadır. Bu tahmin özel illüstrasyon ve pazarlama üretimini sınırlı kabul eder.",
    )

    doc.add_heading("14.1 MVP dışında", level=2)
    add_list(
        doc,
        [
            "PvP, klanlar, oyuncular arası pazar ve liderlik tablosu.",
            "Canlı MMO dünyası, ittifaklar, sezon sıfırlaması, diğer mini oyunlar ve 3D dünya haritası.",
            "Telegram/Discord botu, mobil uygulama ve push bildirimi.",
            "Season Pass, ödeme sistemi ve gerçek para ekonomisi.",
            "Kullanıcı üretimli görev veya terminal komutu paylaşımı.",
            "Çok bölge yüksek erişilebilirlik ve mikroservis ayrıştırması.",
        ],
    )

    doc.add_heading("14.2 İlk dört sprint", level=2)
    add_table(
        doc,
        ["Sprint", "Hedef", "Somut teslimat"],
        [
            ("S1", "Repo ve ürün kabuğu", "React/Nest monorepo, CI, Docker, tasarım tokenları, auth iskeleti"),
            ("S2", "Ekonomi çekirdeği", "Beş kaynak ledger'ı, lazy accrual, enerji önceliği, tesis seviyesi 1-3"),
            ("S3", "Yaşayan dünya", "NPC çevrimi, pazar endeksi, sözleşme seçimi ve deterministik dünya olayı"),
            ("S4", "İlk birleşik dilim", "Dünya olayı -> hazırlık -> nötr operasyon çözümü -> pazar ve ilişki sonucu"),
        ],
        [1200, 3000, 5160],
        font_size=9.0,
    )


def add_acceptance_criteria(doc: Document) -> None:
    doc.add_heading("15. MVP kabul kriterleri ve karar kapıları", level=1)
    add_table(
        doc,
        ["Alan", "Kabul kriteri"],
        [
            ("Oynanış", "Yeni hesap en az 3 farklı hazırlıkla uçtan uca operasyon tamamlar; rapor bütün çarpanları ve dünya etkisini gösterir."),
            ("Denge", "Eş-kademe başarı oranı %60-85; bir üst kademe riskli fakat mümkün; mini oyun farkı <=10 yüzde puan."),
            ("Ekonomi", "Negatif bakiye/çifte harcama yok; pazar 0.85-1.15 bandında; 2-10 operasyon/gün güç farkı <%20."),
            ("Yaşayan dünya", "NPC çevrimi deterministik tekrar üretilebilir; sözleşme, pazar ve ilişki değişimleri kayıttan açıklanabilir."),
            ("Dayanıklılık", "Worker/Redis yeniden başlatıldığında finish_at uzlaşması görev ve yükseltmeleri doğru tamamlar."),
            ("Güvenlik", "ASVS 5.0 L2 kapsam matrisi; kritik/yüksek bulgu yok; istemci sonucu değiştiremiyor."),
            ("Erişilebilirlik", "Temel akış yalnız klavye ile; reduced motion; renk bağımsız durum; kritik WCAG 2.2 AA kontrolleri."),
            ("Performans", "1440x900 orta sınıf cihazda hedef 60 FPS; API p95 <250 ms; durum olayı p95 <500 ms."),
            ("Ürün", "İlk görev >=%80, tamamlama >=%70, ikinci görev >=%60, strateji-operasyon bağı >=%75, adalet >=4/5."),
        ],
        [1900, 7460],
        font_size=9.0,
    )
    add_callout(
        doc,
        "Go/no-go",
        "İkinci görevi gönüllü başlatma ve algılanan adalet hedefleri karşılanmıyorsa içerik hacmi artırılmaz. Önce çekirdek döngü ve sonuç açıklaması yeniden tasarlanır.",
        "caution",
    )


def add_backlog(doc: Document) -> None:
    doc.add_heading("16. Uygulanabilir ürün backlog'u", level=1)
    add_table(
        doc,
        ["Epik", "P0 kapsam", "Done tanımı"],
        [
            ("E1 Kimlik ve profil", "Kayıt/giriş, profil oluşturma, güvenli oturum", "Oturum iptali bütün socketleri kapatır"),
            ("E2 Kaynak ledger", "Enerji/İşlem/Bileşen/Sermaye/Uzmanlık, delil ve ısı", "Her değişim atomik ledger kaydı"),
            ("E3 Tesis ve güç", "5 tesis, lazy accrual, enerji önceliği, ücretsiz kuyruk", "Süre, maliyet ve güç planı sürümlü"),
            ("E4 Dünya simülasyonu", "Asteria çevrimi, kriz, NPC kapasitesi ve ilişki", "Aynı seed ve sürüm aynı çevrimi üretir"),
            ("E5 NPC pazarı", "Arz-talep, 0.85-1.15 endeks, delil satışı", "120 günlük stres testi sınırlar içinde"),
            ("E6 Sözleşme", "Hikâye/pazar sözleşmesi, teklif, arşiv", "Kayıp sınırlı ve ilerlemeyi kilitlemiyor"),
            ("E7 Hedef ve plan", "Kademe, seed, istihbarat, yükleme, tahmin bandı", "Hedef planlamadan sonra değişmez"),
            ("E8 Resolver", "Başarı, tespit, ödül, ısı ve dünya sonucu", "Seed ile deterministik tekrar"),
            ("E9 Node Routing", "Graf üretimi, girişler, skor", "Her seed geçerli yol ve erişilebilir kontrol"),
            ("E10 Rapor", "Neden ağacı, kaynak/pazar/ilişki farkı, replay", "Oyuncu tüm çarpanları görebilir"),
            ("E11 İçerik", "3 NPC, 3 görev, 12 hedef, 1 bölge", "Schema doğrulamalı ve sürümlü"),
            ("E12 Telemetri", "Olaylar, dashboard, consent", "PII minimum ve schema version"),
        ],
        [1900, 4100, 3360],
        font_size=8.7,
    )
    doc.add_heading("16.1 İlk uygulanacak ince dilim", level=2)
    add_list(
        doc,
        [
            "Sahte kullanıcıyla Asteria dünya ekranı, beş kaynak ve enerji önceliği.",
            "Tek NPC sözleşmesi, sabit Kademe 1 hedef ve iki yükleme seçeneği.",
            "Sunucu tarafında başarı/tespit önizlemesi ve resolver.",
            "Node Routing yerine ilk hafta nötr skor düğmesi; sonra gri kutu mini oyun.",
            "Sonuç raporunda A, D, başarı, tespit, ödül, ısı, pazar ve NPC ilişkisi.",
            "Bu akış eğlenceli bulunmadan araştırma ağacı, ikinci bölge veya PvP başlanmaz.",
        ],
        ordered=True,
    )


def add_appendices(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Ek A. Formül sabitleri", level=1)
    add_table(
        doc,
        ["Sabit", "Değer", "Gerekçe / ayar yönü"],
        [
            ("Başarı tabanı", "0.08", "Mutlak kilidi önler; daha adil istenirse yükseltilir"),
            ("Başarı genliği", "0.84", "Tavanı 0.92 yapar"),
            ("Başarı ölçeği", "11 puan", "Daha küçük değer daha keskin güç farkı"),
            ("İstihbarat", "0.18 x ln(1+i)", "Azalan marjinal fayda"),
            ("Mini oyun", "[-0.35,+0.40] logit", "Hazırlığı geçersiz kılmayan beceri payı"),
            ("Tespit tabanı/genliği", "0.03 / 0.55", "Tespit ceza ama kesin kader değil"),
            ("Isı katkısı", "0.25 x Isı", "100 ısıda +25 tespit puanı"),
            ("Isı decay", "4/saat", "Zorunlu bekleme olmadan toparlanma"),
            ("Tespit ödülü", "x0.72", "Gizliliği değerli kılar"),
            ("Başarısız delil", "x0.25", "Başarısız oturumu boş bırakmaz"),
            ("Tekrar tabanı", "0.55", "Farm azaltılır, görev değersizleşmez"),
            ("Ödül jitter", "0.92-1.08", "Küçük değişkenlik; sonuç açıklanabilir kalır"),
            ("Üretim büyümesi", "1.24", "Beş tesis için ortak başlangıç eğrisi"),
            ("Sermaye/bileşen büyümesi", "1.55 / 1.48", "Uzun vadede yatırım ağırlığı"),
            ("Süre büyümesi", "1.50; tavan 360 dk", "İlk hız, ileri aşamada sınırlı bekleme"),
            ("Çevrimdışı üst sınır", "24-36 saat", "Günde çoklu giriş baskısını kaldırır"),
            ("Aktif ekonomi tavanı", "%20", "Oturum süresinin ham güce dönüşmesini sınırlar"),
            ("Pazar koridoru", "0.85-1.15", "NPC ekonomisinde enflasyon/çöküş sınırı"),
            ("NPC adaptasyon tavanı", "+10 savunma", "Çeşitlilik sağlar, eski araçları geçersiz kılmaz"),
            ("XP eğrisi", "60 x L^1.35", "Erken hızlı, sonra kontrollü ilerleme"),
        ],
        [2400, 2200, 4760],
        font_size=8.7,
    )

    doc.add_heading("Ek B. Örnek operasyon hesabı", level=1)
    add_body(
        doc,
        "Operatör profili Kademe 2 kurumsal ağa giriyor: A=47.24, araç=6, D=37.24, istihbarat=2, mini oyun skoru=58, ısı=12. Isı 40 altında olduğu için başarı cezası yoktur.",
    )
    add_formula(doc, "z = (47.24 + 6 - 37.24)/11 + 0.18xln(3) + (58-50)/125 = 1.716")
    add_formula(doc, "P_başarı = 0.08 + 0.84xsigmoid(1.716) = yaklaşık %79")
    add_formula(doc, "skor_gürültüsü = 6 - 0.12x58 = -0.96")
    add_formula(doc, "q = (36 + 8 - 0.96 + 0.25x12 - 52 - 4)/12 = -0.831")
    add_formula(doc, "P_tespit = 0.03 + 0.55xsigmoid(-0.831) = yaklaşık %20")
    add_body(
        doc,
        "100.000 koşuluk model aynı profil için yaklaşık %79 başarı, %20 tespit, deneme başına 265 sermaye ve 8.3 beklenen ısı üretir. Başarısızlıkta taban delilin %25'i analiz verisi olarak korunur. Küçük farklar skor dağılımı ve ödül jitter'ından gelir.",
    )

    doc.add_heading("Ek C. Araştırma ve kaynakça", level=1)
    add_body(
        doc,
        "Erişim tarihi: 9 Ağustos 2026. Teknik seçimlerde resmî dokümantasyon; oyun motivasyonu kararlarında birincil akademik çalışmalar kullanılmıştır.",
    )
    sources = [
        "[R1] NIST, The Cybersecurity Framework (CSF) 2.0, NIST CSWP 29 (2024). https://doi.org/10.6028/NIST.CSWP.29",
        "[R2] MITRE ATT&CK, Enterprise Tactics; current v19.2 update dated 6 August 2026. https://attack.mitre.org/tactics/",
        "[R3] NIST SP 800-61r3, Incident Response Recommendations and Considerations for Cybersecurity Risk Management (2025). https://doi.org/10.6028/NIST.SP.800-61r3",
        "[R4] CISA, Cross-Sector Cybersecurity Performance Goals. https://www.cisa.gov/cybersecurity-performance-goals",
        "[R5] Przybylski, Rigby & Ryan, A Motivational Model of Video Game Engagement, Review of General Psychology 14(2), 2010. https://selfdeterminationtheory.org/SDT/documents/2010_PrzybylskiRigbyRyan_ROGP.pdf",
        "[R6] Ryan & Deci, Self-Determination Theory and the Facilitation of Intrinsic Motivation, American Psychologist 55(1), 2000. https://www.selfdeterminationtheory.org/SDT/documents/2000_RyanDeci_SDT.pdf",
        "[R7] Hunicke, LeBlanc & Zubek, MDA: A Formal Approach to Game Design and Game Research, 2004. https://www.cs.northwestern.edu/~hunicke/MDA.pdf",
        "[R8] OWASP, Application Security Verification Standard 5.0.0. https://owasp.org/www-project-application-security-verification-standard/",
        "[R9] OWASP Cheat Sheet Series, WebSocket Security. https://cheatsheetseries.owasp.org/cheatsheets/WebSocket_Security_Cheat_Sheet.html",
        "[R10] NIST SP 800-218, Secure Software Development Framework 1.1. https://doi.org/10.6028/NIST.SP.800-218",
        "[R11] React Documentation, Managing State and TypeScript. https://react.dev/learn/managing-state",
        "[R12] PixiJS 8 Guides, Performance Tips and Renderers. https://pixijs.com/8.x/guides/concepts/performance-tips",
        "[R13] NestJS Documentation, Queues/BullMQ. https://docs.nestjs.com/techniques/queues",
        "[R14] PostgreSQL 18 Documentation, Transaction Processing. https://www.postgresql.org/docs/current/transactions.html",
        "[R15] Redis Documentation, Streams and Data Types. https://redis.io/docs/latest/develop/data-types/streams/",
        "[R16] W3C, Web Content Accessibility Guidelines (WCAG) 2.2. https://www.w3.org/TR/WCAG22/",
        "[R17] Hades' Star, Frequently Asked Questions; persistent world and removal of offline resource theft pressure. https://hadesstar.com/faq.html",
        "[R18] Hades' Star, Official Features; play at your own pace and no offline resource theft. https://hadesstar.com/en/index.html",
        "[R19] Prosperous Universe, About; browser-based company management, production chains, persistent economy and play-at-your-own-pace design. https://prosperousuniverse.com.s3.amazonaws.com/about/index.html",
        "[R20] Fair Play Alliance & ADL, Disruption and Harms in Online Gaming Framework (2020). https://www.adl.org/resources/report/disruption-and-harms-online-gaming-framework",
        "[R21] Microsoft Research, TrueSkill: A Bayesian Skill Rating System (2006); future optional PvP matchmaking research. https://www.microsoft.com/en-us/research/?p=154591",
    ]
    add_list(doc, sources, ordered=False)

    doc.add_heading("Ek D. Tasarım karar kaydı", level=1)
    add_table(
        doc,
        ["Karar", "Seçim", "Yeniden değerlendirme tetikleyicisi"],
        [
            ("PvE önce", "NPC hedefler ve kalıcı profil", "Çekirdek döngü go/no-go hedeflerini geçtikten sonra PvP araştırılır"),
            ("Birleşik ürün", "Strateji dünyası + Ghost Grid operasyonu", "Oyuncular iki katmanı bağlantısız algılarsa kapsam yeniden kesilir"),
            ("Güvenli merkez", "Kalıcı tesisler saldırıya kapalı", "Değişmez; gelecek PvP yalnız dış fırsatlarda"),
            ("Beş ana kaynak", "Enerji, İşlem, Bileşen, Sermaye, Uzmanlık", "İlk testte kaynak rolleri karışıyorsa azaltılır"),
            ("NPC pazarı", "Kontrollü 0.85-1.15 koridoru", "Oyuncu tercihleri anlamsız kalırsa arz-talep modeli derinleştirilir"),
            ("2D harita", "PixiJS WebGL topoloji", "Okunabilirlik kanıtlandıktan sonra görsel derinlik artırılır"),
            ("Tek mini oyun", "Node Routing", "İkinci mini oyun ancak ilk oyun tekrar yorgunluğu yaratırsa"),
            ("Modüler monolit", "Tek API + worker", "Ölçülmüş ölçek veya ekip sahipliği ihtiyacı"),
            ("Şeffaf DDA", "Açık yardım ve hedef önerisi", "Oyuncular sistemi sömürüyor veya adaletsiz buluyorsa"),
            ("Kozmetik gelir", "Prototipte ödeme yok", "Ürün eğlencesi ve retention doğrulandıktan sonra"),
        ],
        [2050, 3300, 4010],
        font_size=8.9,
    )


def structural_audit(doc: Document) -> list[str]:
    issues: list[str] = []
    section = doc.sections[0]
    if round(section.page_width.inches, 2) != 8.5 or round(section.page_height.inches, 2) != 11.0:
        issues.append("Page size is not US Letter")
    for side in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin):
        if round(side.inches, 2) != 1.0:
            issues.append("Margin differs from 1 inch")
    for idx, table in enumerate(doc.tables):
        tblw = table._tbl.tblPr.find(qn("w:tblW"))
        if tblw is None or tblw.get(qn("w:type")) != "dxa" or int(tblw.get(qn("w:w"))) != CONTENT_WIDTH_DXA:
            issues.append(f"Table {idx} width invalid")
        tblind = table._tbl.tblPr.find(qn("w:tblInd"))
        if tblind is None or int(tblind.get(qn("w:w"))) != TABLE_INDENT_DXA:
            issues.append(f"Table {idx} indent invalid")
        expected = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid.findall(qn("w:gridCol"))]
        if sum(expected) != CONTENT_WIDTH_DXA:
            issues.append(f"Table {idx} grid sum invalid")
        for row in table.rows:
            for cell_idx, cell in enumerate(row.cells):
                tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                if tcw is None or tcw.get(qn("w:type")) != "dxa" or int(tcw.get(qn("w:w"))) != expected[cell_idx]:
                    issues.append(f"Table {idx} cell width mismatch")
                    break
    return sorted(set(issues))


def build() -> Path:
    if not BALANCE_PATH.exists():
        raise FileNotFoundError("Run tools/balance_model.py first")
    doc = Document()
    configure_document(doc)
    setup_running_header_footer(doc)
    doc.core_properties.title = "NEXUS 2040: Ghost Grid - GDD v1.1"
    doc.core_properties.subject = "Kalıcı teknoloji stratejisi ve PvE siber operasyon simülasyonu; denge matematiği ve geliştirme planı"
    doc.core_properties.author = "Proje ekibi"
    doc.core_properties.keywords = "GDD, PvE, strategy, simulation, economy, cybersecurity, game balance, NEXUS 2040, Ghost Grid"
    doc.core_properties.comments = "Birleşik araştırma temelli üretim spesifikasyonu; 9 Ağustos 2026"

    add_cover(doc)
    add_toc(doc)
    doc.add_page_break()
    add_executive_summary(doc)
    add_research_foundation(doc)
    add_vision_and_experience(doc)
    add_core_loop(doc)
    add_operation_system(doc)
    add_missions_and_ai(doc)
    add_minigame(doc)
    add_economy(doc)
    add_retention(doc)
    add_ui_ux(doc)
    add_technical_architecture(doc)
    add_security_and_ethics(doc)
    add_telemetry_and_balance(doc)
    add_roadmap(doc)
    add_acceptance_criteria(doc)
    add_backlog(doc)
    add_appendices(doc)

    issues = structural_audit(doc)
    if issues:
        raise RuntimeError("Structural audit failed: " + "; ".join(issues))
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
    print("Structural preset audit: PASS")
    return OUTPUT


if __name__ == "__main__":
    build()
